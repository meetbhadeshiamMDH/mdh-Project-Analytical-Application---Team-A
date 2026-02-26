"""
Generates the Berlin Bike Theft Dashboard Project Report (.docx)
Run:  python generate_report.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, sys

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Helper functions ─────────────────────────────────────────────────────────

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    elif level == 2:
        p.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def para(text, bold=False, italic=False, size=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    return p

def code_block(code_text):
    """Add a shaded code block."""
    p = doc.add_paragraph()
    p.style = doc.styles['No Spacing']
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    # Grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    # Border
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '4')
        bdr.set(qn('w:space'), '4')
        bdr.set(qn('w:color'), 'AAAAAA')
        pBdr.append(bdr)
    pPr.append(pBdr)
    return p

def insert_image(path, width=Inches(5.5), caption_text=None):
    """Insert image if found, else insert placeholder note."""
    if path and os.path.exists(path):
        doc.add_picture(path, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[ Screenshot: {caption_text or os.path.basename(str(path))} ]")
        run.font.color.rgb = RGBColor(0x99,0x99,0x99)
        run.font.italic = True
    if caption_text:
        cp = doc.add_paragraph(caption_text)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.italic = True
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def divider():
    doc.add_paragraph("─" * 80)

# Screenshot paths — user should place screenshots here or update paths
SCREENSHOT_DIR = os.path.dirname(os.path.abspath(__file__))

def ss(name):
    """Try to find screenshot by filename variants."""
    for ext in ['.png', '.jpg', '.jpeg', '.PNG', '.JPG']:
        p = os.path.join(SCREENSHOT_DIR, name + ext)
        if os.path.exists(p):
            return p
    return None  # will fall back to placeholder

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Berlin Bike Theft Analytical Dashboard")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

doc.add_paragraph()
t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("IT Project — Report Documentation")
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run("Medizinische Hochschule Hannover\nTeam A  •  2024 / 2025")
r3.font.size = Pt(12)
r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
insert_image(ss("dashboard_main"), width=Inches(5.8),
             caption_text="Figure 1: Security Intelligence Dashboard — Main Overview")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════════════════
heading("Table of Contents", 1)
toc_items = [
    ("1.", "Introduction"),
    ("2.", "System Architecture"),
    ("3.", "Technology Stack"),
    ("4.", "User Requirements"),
    ("5.", "Data Pipeline — Ingestion, Storage & Cleaning"),
    ("6.", "Dashboard Features"),
    ("  6.1", "Main Dashboard — KPI Cards"),
    ("  6.2", "Statistics & Filtering View"),
    ("  6.3", "Hourly Theft Trends Chart"),
    ("  6.4", "Weekly Theft Trends Chart"),
    ("  6.5", "Monthly Theft Trends Chart"),
    ("  6.6", "Yearly Theft Trends Chart"),
    ("  6.7", "Financial Impact Trend Chart"),
    ("  6.8", "Geospatial Intelligence — Heatmap"),
    ("  6.9", "Financial Comparison View"),
    ("7.", "Analytical Insights from the Data"),
    ("8.", "Key Learnings"),
    ("9.", "Conclusion"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.add_run(f"{num}   {title}").font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
heading("1. Introduction")
para(
    "Bicycle theft is a persistent and significant urban problem in Berlin. "
    "With tens of thousands of incidents reported each year, understanding the "
    "patterns behind these thefts is essential for city planners, law enforcement, "
    "and citizens. This project was undertaken as part of an IT course, with the "
    "goal of building a full-stack, data-driven analytical application that transforms "
    "raw police incident data into actionable visual insights."
)
para(
    "The 'Berlin Bike Theft Analytical Dashboard' (BikeGuard v1.0) provides an "
    "interactive web interface for exploring theft incidents across multiple dimensions: "
    "time of day, day of week, month, year, geographic district, bicycle type, and "
    "financial damage. The application was built using modern web technologies and "
    "follows a clean client-server architecture."
)

# ════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("2. System Architecture")
para(
    "The application follows a two-tier client-server architecture with a clear "
    "separation of concerns between the frontend and backend layers."
)
heading("Backend (Server)", 2)
bullet("Language: Python 3.x")
bullet("Framework: Flask — lightweight REST API server")
bullet("Data Loading: Excel file read at application startup via Pandas + OpenPyXL")
bullet("API: RESTful JSON endpoints served at http://localhost:5000")
bullet("CORS: Flask-CORS enables cross-origin requests from the React frontend")

heading("Frontend (Client)", 2)
bullet("Framework: React 18 with Vite build tooling")
bullet("Language: JavaScript (ES6+) / JSX")
bullet("Charts: Recharts library (BarChart, LineChart, ResponsiveContainer)")
bullet("Map: Leaflet.js + react-leaflet (interactive choropleth heatmap)")
bullet("HTTP Client: Axios for REST API calls")
bullet("Styling: CSS Custom Properties (design tokens) with dark-mode theme")

heading("Data Flow", 2)
para(
    "Excel File → Python (Pandas cleaning) → Flask REST API → Axios HTTP → "
    "React State → Recharts / Leaflet Rendering → User Browser"
)

code_block(
"""# api.py — Application startup: data loaded once into memory
raw_df = load_data(DATA_FILE)          # reads .xlsx via openpyxl
df     = clean_data(raw_df)            # normalises dates, types, LOR codes
app.run(debug=True, port=5000)         # Flask serves all /api/stats/* routes"""
)

# ════════════════════════════════════════════════════════════════════════════
# 3. TECHNOLOGY STACK
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("3. Technology Stack")

heading("3.1 Python Libraries (Backend)", 2)

libs = [
    ("Flask",       "REST API framework. Defines all /api/stats/* HTTP endpoints. "
                    "Handles routing, JSON serialisation, and request handling."),
    ("Flask-CORS",  "Enables Cross-Origin Resource Sharing so the React app "
                    "(port 5173) can make requests to the Flask server (port 5000) "
                    "without browser security blocks."),
    ("Pandas",      "Core data manipulation library. Used for loading the Excel file, "
                    "cleaning columns, parsing dates, aggregating theft counts by hour/"
                    "week/month/year, and grouping financial damage by LOR district."),
    ("NumPy",       "Numerical computation support. Used for type conversions (float/int) "
                    "during data cleaning to handle NaN values gracefully."),
    ("OpenPyXL",    "Excel engine used by Pandas to read the .xlsx data file. Enables "
                    "Pandas to open and parse multi-sheet Excel workbooks."),
]
for lib, desc in libs:
    p = doc.add_paragraph()
    p.add_run(f"{lib}:  ").bold = True
    p.add_run(desc)
    p.style = doc.styles['No Spacing']
    doc.add_paragraph()

heading("3.2 JavaScript / React Libraries (Frontend)", 2)
jslibs = [
    ("React 18",       "Component-based UI framework. Each chart and view is an isolated component."),
    ("Vite",           "Fast build tool and dev server with Hot Module Replacement (HMR)."),
    ("Recharts",       "Chart library built on D3. Provides BarChart, LineChart, XAxis, YAxis, "
                       "Tooltip, ResponsiveContainer used in all 5 statistical charts."),
    ("Axios",          "Promise-based HTTP client used to call all Flask REST endpoints."),
    ("Leaflet.js",     "Interactive map library used for the Geospatial heatmap view."),
    ("react-leaflet",  "React wrapper for Leaflet. Provides MapContainer, TileLayer, GeoJSON components."),
]
for lib, desc in jslibs:
    p = doc.add_paragraph()
    p.add_run(f"{lib}:  ").bold = True
    p.add_run(desc)
    p.style = doc.styles['No Spacing']
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 4. USER REQUIREMENTS
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("4. User Requirements")
para(
    "The requirements for this dashboard were gathered through a stakeholder interview. "
    "The primary user (a data analyst / criminologist) expressed clear needs around "
    "spatial and temporal analysis of bike theft data. Key requirements are summarised below."
)

para('"Absolute numbers distort the picture. The prevalence rate is key for spotting high-risk areas."',
     italic=True)
doc.add_paragraph()

# Requirements table
from docx.oxml.ns import qn as _qn
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

hdr = table.rows[0].cells
hdr[0].text = "Requirement"
hdr[1].text = "Details"
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
    shd = OxmlElement('w:shd')
    shd.set(_qn('w:val'), 'clear')
    shd.set(_qn('w:color'), 'auto')
    shd.set(_qn('w:fill'), 'BDD7EE')
    cell._tc.get_or_add_tcPr().append(shd)

rows_data = [
    ("Primary Need",       "Time trends — identify patterns and regional hotspots"),
    ("Key Metric",         "Prevalence rate (cases per 100k inhabitants) — essential because "
                           "absolute numbers are misleading due to population differences"),
    ("Example Question",   "Was August 2025 more critical than August 2024 or 2023?"),
    ("Time Series View",   "Line charts with 3-year overlay (2023–2025) with optional "
                           "annotations for interventions (e.g., 'secure parking pilot launched')"),
    ("Spatial Level",      "District regions (BZR level). PLR is too granular; "
                           "BZR is a good compromise for clear communication"),
    ("Metrics per BZR",    "Case count, damage sum, average damage — aggregated monthly or yearly"),
    ("Bicycle Type",       "INCREASING interest — high-value e-bikes stolen more often. "
                           "Desired: table by year showing type × count, prevalence, damage"),
    ("Top Cases View",     "Sort incidents by damage to identify organized theft groups. "
                           "Details needed: PLR, time, bike type, comparison to historical highs"),
    ("Export Need",        "Excel export for further analysis by the user"),
    ("Comparative Tables", "All BZR regions sortable by theft frequency and financial damage"),
]
for req, detail in rows_data:
    row = table.add_row().cells
    row[0].text = req
    row[1].text = detail

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 5. DATA PIPELINE
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("5. Data Pipeline — Ingestion, Storage & Cleaning")

heading("5.1 Data Source", 2)
para(
    "The raw data comes from the Berlin Police Department's open-data portal, "
    "provided as an Excel file (bike_thefts_berlin.xlsx). The dataset contains "
    "approximately 59,707 individual theft incident records with the following key fields:"
)
for field in ["Start date / End date", "Start hour / End hour",
              "LOR district code (8-digit PLR_ID)", "Type of bicycle",
              "Financial damage (€)", "Created on (report date)"]:
    bullet(field)

heading("5.2 Data Ingestion", 2)
para(
    "The data ingestion module (data_ingestion.py) reads the Excel file using "
    "Pandas with the OpenPyXL engine. The file is loaded once at application "
    "startup for performance efficiency:"
)
code_block(
"""# data_ingestion.py
import pandas as pd

def load_data(filepath):
    \"\"\"Load the Excel bike theft dataset into a DataFrame.\"\"\"
    try:
        df = pd.read_excel(filepath, engine='openpyxl')
        print(f"Loaded {len(df)} records from {filepath}")
        return df
    except FileNotFoundError:
        print(f"ERROR: Data file not found at {filepath}")
        return pd.DataFrame()"""
)

heading("5.3 Data Cleaning", 2)
para(
    "The cleaning module (data_processing.py) standardises the raw data for "
    "consistent analysis. Key cleaning steps include:"
)
bullet("Date parsing: Convert Start date, End date, Created on to datetime objects")
bullet("Numeric conversion: Start hour, End hour, Financial damage cast to numeric (NaN on error)")
bullet("LOR standardisation: LOR codes padded with leading zeros to 8 digits (e.g. '9100101' → '09100101')")
bullet("Null handling: errors='coerce' used throughout to silently ignore malformed values")

code_block(
"""# data_processing.py — clean_data()
def clean_data(df):
    if 'Start date' in df.columns:
        df['Start date'] = pd.to_datetime(df['Start date'], errors='coerce')

    numeric_cols = ['Start hour', 'End hour', 'financial damage']
    for col in numeric_cols:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors='coerce')

    # Standardise LOR to 8-digit string codes
    if 'LOR' in df.columns:
        df['LOR'] = df['LOR'].astype(str).str.replace(r'\\.0$', '', regex=True)
        df['LOR'] = df['LOR'].apply(
            lambda x: x.zfill(8) if x not in ('nan', 'None') else x
        )
    return df"""
)

heading("5.4 Data Aggregation & Analytics", 2)
para(
    "Once cleaned, the data is aggregated by multiple dimensions using Pandas "
    "groupby and value_counts operations. Each aggregation is exposed via a "
    "dedicated Flask API endpoint:"
)
agg_table = doc.add_table(rows=1, cols=3)
agg_table.style = 'Table Grid'
hdr2 = agg_table.rows[0].cells
hdr2[0].text = "Dimension"
hdr2[1].text = "Method"
hdr2[2].text = "API Endpoint"
for cell in hdr2:
    cell.paragraphs[0].runs[0].bold = True

agg_rows = [
    ("Hour of day",     "df['Start hour'].value_counts().sort_index()",     "/api/stats/hourly"),
    ("Day of week",     "dt.day_name().value_counts() + reindex",           "/api/stats/weekly"),
    ("Month",          "dt.month_name().value_counts() + reindex",          "/api/stats/monthly"),
    ("Year",           "dt.year.value_counts().sort_index()",               "/api/stats/yearly"),
    ("Financial loss", "dt.to_period('M') + groupby['financial damage'].sum()", "/api/stats/financial"),
    ("LOR district",   "groupby('LOR').agg(count + damage sum)",            "/api/stats/geospatial"),
    ("Summary KPIs",   "len(df), mean(damage), mode(hour), mode(LOR)",      "/api/stats/summary"),
    ("Filtered count", "POST body filters applied to DataFrame copy",       "/api/stats/filtered"),
]
for d, m, e in agg_rows:
    r = agg_table.add_row().cells
    r[0].text = d; r[1].text = m; r[2].text = e

# ════════════════════════════════════════════════════════════════════════════
# 6. DASHBOARD FEATURES
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("6. Dashboard Features")

# ── 6.1 Main Dashboard ──────────────────────────────────────────────────────
heading("6.1 Main Dashboard — KPI Cards", 2)
para(
    "The main dashboard provides an at-a-glance summary of the most important "
    "metrics from the dataset through four Key Performance Indicator (KPI) cards:"
)
bullet("Total Theft Reports: 59,707 — the total number of recorded incidents")
bullet("Average Loss: €1,224 per incident — mean financial damage across all thefts")
bullet("Peak Activity: 18:00 — the hour of day with the highest theft frequency")
bullet("High-Risk Sector: 09100101 — the LOR district with the most thefts")
insert_image(ss("dashboard_main"), width=Inches(5.5),
             caption_text="Figure 1: Security Intelligence Dashboard — KPI Overview")

# ── 6.2 Statistics View ─────────────────────────────────────────────────────
doc.add_page_break()
heading("6.2 Statistics & Advanced Filtering View", 2)
para(
    "The Statistics view provides advanced filtering capabilities allowing users "
    "to narrow down the dataset by bicycle type and year. Results update "
    "dynamically showing the filtered total case count and average financial damage."
)
bullet("Filter by bicycle type: Bicycle, Cargo bike, Children's bicycle, E-bike, Mountain bike, etc.")
bullet("Filter by year: 2023, 2024, 2025")
bullet("Dynamic result: Total cases and average damage update instantly via /api/stats/filtered")
insert_image(ss("statistics_view"), width=Inches(5.5),
             caption_text="Figure 2: Statistics — Advanced Filtering Panel")

# ── 6.3 Hourly Chart ────────────────────────────────────────────────────────
doc.add_page_break()
heading("6.3 Hourly Theft Trends Chart", 2)
para(
    "The hourly chart visualises theft frequency across each hour of the day (0–23). "
    "This helps identify the most dangerous times for bike theft."
)
bullet("Chart Type: Bar Chart (Recharts BarChart)")
bullet("X-Axis: Hour of day (0 to 23)")
bullet("Y-Axis: Number of theft incidents")
bullet("Key Insight: Peak theft occurs at 17:00–18:00 (commuter evening hours) with ~6,500+ incidents")
bullet("Lowest Risk: 02:00–05:00 (early morning, fewest incidents)")
bullet("API: GET /api/stats/hourly → { '0': 800, '1': 220, ..., '17': 6500, ... }")

insert_image(ss("hourly_chart"), width=Inches(4.0),
             caption_text="Figure 3: Hourly Theft Trends — Bar Chart")

code_block(
"""// HourlyChart.jsx — key code snippet
import { BarChart, Bar, XAxis, YAxis, Tooltip, ResponsiveContainer } from 'recharts';

const response = await axios.get('http://localhost:5000/api/stats/hourly');
const formattedData = Object.entries(response.data)
    .map(([hour, count]) => ({ hour: parseInt(hour), count }))
    .sort((a, b) => a.hour - b.hour);

<BarChart data={formattedData}>
    <XAxis dataKey="hour" />
    <YAxis />
    <Tooltip />
    <Bar dataKey="count" fill="var(--accent-primary)" radius={[4, 4, 0, 0]} />
</BarChart>"""
)

# ── 6.4 Weekly Chart ────────────────────────────────────────────────────────
doc.add_page_break()
heading("6.4 Weekly Theft Trends Chart", 2)
para(
    "The weekly chart shows theft distribution across the seven days of the week, "
    "revealing which days are most prone to bike theft."
)
bullet("Chart Type: Bar Chart (Recharts BarChart)")
bullet("X-Axis: Day of week (Monday → Sunday)")
bullet("Y-Axis: Number of theft incidents")
bullet("Key Insight: Monday to Friday show consistently high thefts (~8,500–9,500); Sunday lowest (~6,500)")
bullet("Pattern: Weekday thefts correlate with commuter activity; weekend slight decline")
bullet("API: GET /api/stats/weekly → { 'Monday': 8900, 'Tuesday': 8850, ... }")

insert_image(ss("weekly_chart"), width=Inches(4.0),
             caption_text="Figure 4: Weekly Theft Trends — Bar Chart")

code_block(
"""// WeeklyChart.jsx — key code snippet
const weekOrder = ['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday'];
const formattedData = Object.entries(response.data)
    .map(([day, count]) => ({ day, count }))
    .sort((a, b) => weekOrder.indexOf(a.day) - weekOrder.indexOf(b.day));

<Bar dataKey="count" fill="var(--accent-success)" radius={[4, 4, 0, 0]} />"""
)

# ── 6.5 Monthly Chart ───────────────────────────────────────────────────────
doc.add_page_break()
heading("6.5 Monthly Theft Trends Chart", 2)
para(
    "The monthly line chart reveals seasonal patterns in bike theft across the calendar year."
)
bullet("Chart Type: Line Chart (Recharts LineChart)")
bullet("X-Axis: Month name (January → December)")
bullet("Y-Axis: Total theft incidents for that month")
bullet("Key Insight: Summer peak in June–August (~6,500–7,000 thefts/month). Strong winter dip in December (~2,900)")
bullet("Pattern: Clear seasonal sinusoidal curve — more outdoor activity = more theft opportunity")
bullet("API: GET /api/stats/monthly → { 'January': 3700, 'June': 6600, ... }")

insert_image(ss("monthly_yearly_chart"), width=Inches(5.5),
             caption_text="Figure 5: Monthly Theft Trends (Line) & Yearly Trends (Bar)")

code_block(
"""// MonthlyChart.jsx — key code snippet
import { LineChart, Line, XAxis, YAxis, Tooltip, ResponsiveContainer } from 'recharts';

<LineChart data={formattedData}>
    <XAxis dataKey="month" />
    <YAxis />
    <Tooltip />
    <Line type="monotone" dataKey="count"
          stroke="var(--accent-warning)" strokeWidth={3}
          dot={{ fill: 'var(--accent-warning)', r: 4 }} />
</LineChart>"""
)

# ── 6.6 Yearly Chart ────────────────────────────────────────────────────────
heading("6.6 Yearly Theft Trends Chart", 2)
para(
    "The yearly bar chart tracks long-term theft trends across multiple calendar years."
)
bullet("Chart Type: Bar Chart with SVG gradient fill (Recharts BarChart + linearGradient)")
bullet("X-Axis: Year (2022, 2023, 2024, 2025)")
bullet("Y-Axis: Total incidents in that year")
bullet("Key Insight: Data spans 2022–2025 with ~17,000–21,000 incidents per year")
bullet("Trend: Slight decline visible in 2025 (partial year data)")
bullet("API: GET /api/stats/yearly → { '2022': 21800, '2023': 20900, '2024': 16800 ... }")

code_block(
"""// YearlyChart.jsx — gradient bar example
<defs>
    <linearGradient id="yearlyGradient" x1="0" y1="0" x2="0" y2="1">
        <stop offset="0%" stopColor="var(--accent-secondary)" stopOpacity={0.8} />
        <stop offset="100%" stopColor="var(--accent-secondary)" stopOpacity={0.2} />
    </linearGradient>
</defs>
<Bar dataKey="count" fill="url(#yearlyGradient)" radius={[6, 6, 0, 0]} barSize={40} />"""
)

# ── 6.7 Financial Impact Trend ──────────────────────────────────────────────
doc.add_page_break()
heading("6.7 Financial Impact Trend Chart", 2)
para(
    "The financial impact chart tracks total monetary loss from bike theft across "
    "time, grouped by year-month (e.g. '2023-06'). This is the most important chart "
    "for insurance and policy analysis."
)
bullet("Chart Type: Line Chart (Recharts LineChart)")
bullet("X-Axis: Year-Month label (e.g. 2023-03, 2023-07, ...)")
bullet("Y-Axis: Total financial damage in Euros (€)")
bullet("Range: €1.4M – €2.8M per month across the 2023–2025 period")
bullet("Tooltip: Formatted as €X,XXX,XXX for readability")
bullet("Key Insight: Clear cyclical pattern with peaks in mid-year. Overall declining trend in late 2025")
bullet("API: GET /api/stats/financial → { '2023-01': 1680000, '2023-07': 2800000, ... }")

insert_image(ss("financial_chart"), width=Inches(5.5),
             caption_text="Figure 6: Financial Impact Trend — Monthly Loss in Euros (€)")

code_block(
"""// FinancialChart.jsx — key code snippet
const formattedData = Object.entries(response.data)
    .map(([date, amount]) => ({ date, amount }))
    .sort((a, b) => a.date.localeCompare(b.date));

<Tooltip formatter={(value) => `€${value.toLocaleString()}`} />
<Line type="monotone" dataKey="amount"
      stroke="var(--accent-primary)" strokeWidth={3}
      dot={{ fill: 'var(--accent-primary)', r: 4 }} />

# Backend — data_processing.py
temp_df['Month'] = temp_df['Start date'].dt.to_period('M')
financial_counts = temp_df.groupby('Month')['financial damage'].sum().sort_index()"""
)

# ── 6.8 Geodata Heatmap ─────────────────────────────────────────────────────
doc.add_page_break()
heading("6.8 Geospatial Intelligence — District Heatmap", 2)
para(
    "The Geodata view overlays theft statistics onto an interactive map of Berlin "
    "using LOR district boundaries (GeoJSON). Each district is coloured based on "
    "its theft density — creating a choropleth heatmap."
)
bullet("Library: Leaflet.js + react-leaflet (MapContainer, TileLayer, GeoJSON)")
bullet("Base Map: CARTO dark tile layer for a sleek dark-mode appearance")
bullet("Color Scale: Green (<50 thefts) → Yellow → Amber → Orange → Red (200+ thefts)")
bullet("Interactivity: Hover tooltip shows district name, LOR ID, theft count, and total damage")
bullet("API: GET /api/stats/geospatial → { '09100101': { count: 285, damage: 358000 }, ... }")
bullet("GeoJSON: Berlin LOR boundary file (berlin_lor.json) served as static asset")

insert_image(ss("geodata_heatmap"), width=Inches(5.5),
             caption_text="Figure 7: Geospatial Intelligence — Berlin Bike Theft Heatmap by LOR District")

code_block(
"""// GeodataDashboard.jsx — colour scaling + tooltip
const getColor = (count) => {
    if (!count)    return '#1a202c';   // no data
    return count > 200 ? '#ef4444' :   // Red   — Critical
           count > 150 ? '#f97316' :   // Orange — High
           count > 100 ? '#f59e0b' :   // Amber  — Moderate
           count > 50  ? '#eab308' :   // Yellow — Low-Moderate
                         '#22c55e';    // Green  — Low
};

layer.bindTooltip(
    `<strong>${lorName}</strong><br/>
     Thefts: ${data.count.toLocaleString()}<br/>
     Damage: €${Math.round(data.damage).toLocaleString()}`,
    { sticky: true }
);

# Backend — api.py
geo_stats = df.groupby('LOR').agg({
    'LOR': 'count', 'financial damage': 'sum'
}).rename(columns={'LOR': 'theft_count', 'financial damage': 'total_damage'})"""
)

# ── 6.9 Financial Comparison View ───────────────────────────────────────────
doc.add_page_break()
heading("6.9 Financial Comparison View (Dual Panel)", 2)
para(
    "The Financial Comparison view is a unique analytical feature that allows "
    "users to compare two independent data segments side by side. Each panel "
    "has its own set of filters so users can compare any two combinations."
)
bullet("Panel A (blue) and Panel B (green) are fully independent")
bullet("Chart Type selector: Financial Loss, Monthly, Weekly, Yearly, Hourly counts")
bullet("Filters per panel: Bike Type, Year, Month")
bullet("Charts update instantly when any filter changes")
bullet("Comparison logic: Uses /api/stats/filtered to calculate a scale ratio, then applies it proportionally to the time-series data")
bullet("Use case: Compare 2023 vs 2024 monthly trends; or Cargo Bike vs Mountain Bike financial losses")

insert_image(ss("financials_comparison"), width=Inches(5.5),
             caption_text="Figure 8: Financial Comparison — Dual Panel with Independent Filters")

code_block(
"""// FinancialComparison.jsx — filter + scale logic
const filterBody = {};
if (bikeType !== 'all') filterBody.bike_types = [bikeType];
if (year     !== 'all') filterBody.years      = [parseInt(year)];
if (month    !== 'all') filterBody.months     = [month];

// Scale factor: filtered count / total count
const filteredCount = filteredResp.data.count || 0;
const totalThefts   = summaryResp.data.total_thefts || 1;
const scale = filteredCount / totalThefts;

// Apply scale to all data points
formatted = rawEntries.map(([label, value]) => ({
    label,
    value: Math.round(value * scale)
}));"""
)

# ════════════════════════════════════════════════════════════════════════════
# 7. ANALYTICAL INSIGHTS
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("7. Analytical Insights from the Data")
para(
    "Based on 59,707 recorded bike theft incidents in Berlin, the dashboard "
    "reveals the following key analytical findings:"
)

insights = [
    ("Temporal Peak",    "The highest theft risk occurs between 17:00–18:00 (evening commute). "
                         "Bike owners leaving work are the most vulnerable."),
    ("Seasonal Pattern", "Summer months (June–August) consistently show the highest theft rates "
                         "— nearly double the winter months. Outdoor activity creates more theft opportunities."),
    ("Day of Week",      "Weekdays show slightly higher theft rates than weekends, "
                         "contradicting the assumption that Saturdays are most dangerous. "
                         "This is driven by commuter bike usage patterns."),
    ("Financial Impact", "Monthly financial damage fluctuates between €1.4M and €2.8M. "
                         "The city loses approximately €20–25M per year to bike theft."),
    ("Geographic Concentration", "The central districts of Berlin (Mitte, Prenzlauer Berg, "
                                 "Friedrichshain) show the darkest red on the heatmap — "
                                 "highest absolute theft counts, driven by high population density."),
    ("Multi-year Trend", "Theft counts have remained relatively stable across 2022–2024 "
                         "at ~17,000–21,000 incidents per year, with a slight downward trend in 2025."),
    ("Average Damage",   "Each theft costs an average of €1,224 — significantly impacting "
                         "lower-income residents who rely on bicycles as primary transport."),
]
for title, detail in insights:
    p = doc.add_paragraph()
    p.add_run(f"{title}: ").bold = True
    p.add_run(detail)

# ════════════════════════════════════════════════════════════════════════════
# 8. KEY LEARNINGS
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("8. Key Learnings")
para(
    "This IT project provided hands-on experience across the full stack of "
    "modern web application development and data analytics. The following key "
    "skills and concepts were applied and reinforced:"
)
learnings = [
    "Full-Stack Development: Building a working web application from scratch — "
    "covering backend API design, data processing, frontend development, and deployment.",

    "REST API Design: Designing clean, well-named API endpoints that return structured "
    "JSON data, following standard HTTP conventions (GET, POST).",

    "Data Engineering with Pandas: Loading, cleaning, and aggregating real-world "
    "police data — handling missing values, type errors, and LOR code standardisation.",

    "Interactive Visualisation: Implementing 5 different chart types with Recharts "
    "(bar, line, gradient bar) and a choropleth map with Leaflet.js.",

    "React Component Architecture: Building reusable, stateful components using "
    "React hooks (useState, useEffect, useCallback) and passing data via props.",

    "Geospatial Analysis: Merging statistical data with GeoJSON boundary files "
    "to create color-coded district maps — a core skill in urban data analytics.",

    "Data Limitations Awareness: Understanding that absolute numbers can be "
    "misleading (the 'prevalence rate problem') — high-density districts naturally "
    "show more thefts in absolute terms.",

    "User-Centered Design: Translating stakeholder requirements (from interviews) "
    "into concrete dashboard features — the comparison view and filter system were "
    "directly driven by user requirements.",

    "Cross-Origin Resource Sharing (CORS): Configuring Flask-CORS to allow the "
    "frontend and backend to communicate across different ports.",

    "Version Control & Collaboration: Managing code through Git with a structured "
    "project repository (mdh-Project-Analytical-Application---Team-A).",
]
for l in learnings:
    bullet(l)

# ════════════════════════════════════════════════════════════════════════════
# 9. CONCLUSION
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("9. Conclusion")
para(
    "The Berlin Bike Theft Analytical Dashboard (BikeGuard v1.0) successfully "
    "transforms raw police incident data into a rich, interactive web application "
    "that serves both casual users and data analysts. The system provides clear "
    "answers to the core questions identified in the stakeholder requirements: "
    "when do thefts happen, where are the hotspots, which bike types are targeted, "
    "and what is the financial impact?"
)
para(
    "The project demonstrates the power of combining Python's data ecosystem "
    "(Pandas, Flask) with modern frontend technologies (React, Recharts, Leaflet) "
    "to deliver real analytical value. With 59,707 incidents analysed and visualised "
    "across 9 different dashboard views, this application provides a solid foundation "
    "for further development — including prevalence rate calculations, BZR-level "
    "aggregation, Excel export functionality, and multi-year overlay line charts "
    "as outlined in the user requirements."
)
para(
    "The team has gained substantial practical experience in data engineering, "
    "API development, interactive visualisation, and user-driven design — skills "
    "that form the foundation of modern data-driven application development."
)

doc.add_paragraph()
divider()
para("Berlin Bike Theft Analytical Dashboard — IT Project Report",
     bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Medizinische Hochschule Hannover  •  Team A  •  2024 / 2025",
     align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ─────────────────────────────────────────────────────────────────────
out_path = r"c:\Users\My PC\Desktop\It-project\mdh-Project-Analytical-Application---Team-A\Berlin_Bike_Theft_Dashboard_Report.docx"
doc.save(out_path)
print(f"[OK] Report saved to: {out_path}")
